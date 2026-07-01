#!/usr/bin/env python3
"""`word-polished-doc-collab` 的参考实现工具集。

定位与作用
----------
这个模块服务 `references/presentation-skills/word-polished-doc-collab/scripts/` 下的
六个入口脚本，负责把共享的数据结构、Markdown 语义、DOCX 构建、预览导出、
Markdown lint 和 DOCX QA 收敛到同一处实现。

设计原则
--------
1. 先锁语义，再锁版式，不让每个 CLI 入口各自发明 block 规则。
2. 先支持最核心的 Markdown -> DOCX -> preview -> QA 闭环，再显式暴露边界。
3. 对复杂 Word 能力正确地失败，不用脚本没报错冒充格式已经锁定。
"""

from __future__ import annotations

from dataclasses import asdict
from dataclasses import dataclass
from pathlib import Path
import json
import re
import shutil
import subprocess
import sys
from typing import Any
from typing import Iterable
from zipfile import ZipFile
import xml.etree.ElementTree as ET

SCRIPT_DIR = Path(__file__).resolve().parent
VENDOR_DIR = SCRIPT_DIR.parent / "vendor"
if VENDOR_DIR.exists():
    sys.path.insert(0, str(VENDOR_DIR))

try:
    from docx import Document
    from docx.document import Document as DocumentObject
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches
    from docx.shared import Pt
    from docx.shared import RGBColor
    DOCX_IMPORT_ERROR: Exception | None = None
except Exception as exc:  # noqa: BLE001
    Document = None
    DocumentObject = Any
    WD_SECTION = None
    WD_STYLE_TYPE = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    qn = None
    Inches = None
    Pt = None
    RGBColor = None
    DOCX_IMPORT_ERROR = exc


W_NAMESPACE = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W_VAL = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
DEFAULT_PDF_DPI = 180
INLINE_BOLD_PATTERN = re.compile(r"(\*\*.*?\*\*)")
IMAGE_PATTERN = re.compile(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)")
TABLE_TITLE_PATTERN = re.compile(r"^(表|Table)\s*\d+(?:[-.]\d+)?(?:\s+|:|$).+", re.IGNORECASE)
FIGURE_TITLE_PATTERN = re.compile(r"^(图|Figure|Exhibit)\s*\d+(?:[-.]\d+)?(?:\s+|:|$).+", re.IGNORECASE)
ORDERED_LIST_PATTERN = re.compile(r"^\d+\.\s+")
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+")
NUMERIC_PATTERN = re.compile(
    r"""
    ^
    [+\-]?
    (?:
        \d{1,3}(?:,\d{3})*|\d+
    )
    (?:\.\d+)?
    %?
    x?
    $
    """,
    re.VERBOSE,
)
ENGLISH_PRESET_PROFILES = {
    "teal_consulting_report",
    "red_private_equity_report",
    "blue_editorial_article",
}
SUPPORTED_FONTS = [
    "宋体",
    "楷体",
    "黑体",
    "Times New Roman",
    "Arial",
    "Georgia",
    "Microsoft YaHei",
    "微软雅黑",
]
STYLE_ID_TO_ROLE = {
    "PSDocTitle": "doc_title",
    "PSSubtitle": "subtitle",
    "PSHeading1": "heading_1",
    "PSHeading2": "heading_2",
    "PSHeading3": "heading_3",
    "PSHeading4": "heading_4",
    "PSLead": "lead",
    "PSBody": "body",
    "PSList": "list",
    "PSTableTitle": "table_title",
    "PSTableCaption": "table_caption",
    "PSFigureTitle": "figure_title",
    "PSFigureNote": "figure_note",
    "PSSourceNote": "source_note",
    "PSFigureBody": "figure_body",
    "TealDocTitle": "doc_title",
    "TealSubtitle": "subtitle",
    "TealHeading1": "heading_1",
    "TealHeading2": "heading_2",
    "TealHeading3": "heading_3",
    "TealHeading4": "heading_4",
    "TealLead": "lead",
    "TealBody": "body",
    "TealList": "list",
    "TealTableTitle": "table_title",
    "TealTableCaption": "table_caption",
    "TealFigureTitle": "figure_title",
    "TealFigureNote": "figure_note",
    "TealSourceNote": "source_note",
    "TealFigureBody": "figure_body",
    "RedDocTitle": "doc_title",
    "RedSubtitle": "subtitle",
    "RedHeading1": "heading_1",
    "RedHeading2": "heading_2",
    "RedHeading3": "heading_3",
    "RedHeading4": "heading_4",
    "RedLead": "lead",
    "RedBody": "body",
    "RedList": "list",
    "RedTableTitle": "table_title",
    "RedTableCaption": "table_caption",
    "RedFigureTitle": "figure_title",
    "RedFigureNote": "figure_note",
    "RedSourceNote": "source_note",
    "RedFigureBody": "figure_body",
    "BlueDocTitle": "doc_title",
    "BlueSubtitle": "subtitle",
    "BlueHeading1": "heading_1",
    "BlueHeading2": "heading_2",
    "BlueHeading3": "heading_3",
    "BlueHeading4": "heading_4",
    "BlueLead": "lead",
    "BlueBody": "body",
    "BlueList": "list",
    "BlueTableTitle": "table_title",
    "BlueTableCaption": "table_caption",
    "BlueFigureTitle": "figure_title",
    "BlueFigureNote": "figure_note",
    "BlueSourceNote": "source_note",
    "BlueFigureBody": "figure_body",
}


@dataclass(frozen=True)
class FontProfile:
    """定义一种中英文字体组合。"""

    chinese: str
    latin: str


@dataclass(frozen=True)
class ParagraphSpec:
    """定义某个 block role 对应的段落样式。"""

    style_name: str
    size_pt: float
    bold: bool
    alignment: str
    line_spacing: float
    space_before_pt: float
    space_after_pt: float
    first_line_indent_pt: float = 0.0
    color_hex: str | None = None
    latin_font: str | None = None
    chinese_font: str | None = None
    keep_with_next: bool = False


@dataclass(frozen=True)
class StyleProfile:
    """定义整份文档的版式契约。"""

    name: str
    font_profile: FontProfile
    role_specs: dict[str, ParagraphSpec]
    table_font_size_pt: float
    dense_table_font_size_pt: float
    table_title_position: str
    figure_title_position: str
    table_caption_position: str
    figure_note_position: str
    source_note_position: str
    page_margin_inch: float
    figure_width_inch: float
    page_width_inch: float
    page_height_inch: float
    body_column_count: int = 1
    column_gap_inch: float = 0.28


@dataclass
class Block:
    """表示一个已经识别好的语义块。"""

    role: str
    text: str = ""
    rows: list[list[str]] | None = None
    image_path: str | None = None


@dataclass
class Issue:
    """表示 lint 或 QA 中的一个问题。"""

    severity: str
    code: str
    message: str


@dataclass
class CheckResult:
    """表示单个 QA 检查项的结果。"""

    passed: bool
    detail: str


def build_style_registry() -> dict[str, StyleProfile]:
    """返回当前脚本支持的全部 style profile。"""

    cn_song_times = StyleProfile(
        name="cn_song_times",
        font_profile=FontProfile(chinese="宋体", latin="Times New Roman"),
        role_specs={
            "doc_title": ParagraphSpec("PSDocTitle", 22.0, True, "center", 1.5, 6.0, 6.0),
            "subtitle": ParagraphSpec("PSSubtitle", 15.0, False, "center", 1.5, 6.0, 6.0),
            "heading_1": ParagraphSpec("PSHeading1", 16.0, True, "left", 1.5, 6.0, 6.0, keep_with_next=True),
            "heading_2": ParagraphSpec("PSHeading2", 15.0, True, "left", 1.5, 6.0, 6.0, keep_with_next=True),
            "heading_3": ParagraphSpec("PSHeading3", 14.0, True, "left", 1.5, 6.0, 6.0, keep_with_next=True),
            "heading_4": ParagraphSpec("PSHeading4", 12.0, True, "left", 1.5, 6.0, 6.0, keep_with_next=True),
            "lead": ParagraphSpec("PSLead", 12.0, True, "justify", 1.5, 0.0, 8.0),
            "body": ParagraphSpec("PSBody", 12.0, False, "justify", 1.5, 6.0, 6.0, 24.0),
            "list": ParagraphSpec("PSList", 12.0, False, "left", 1.5, 0.0, 6.0),
            "table_title": ParagraphSpec("PSTableTitle", 12.0, True, "center", 1.5, 6.0, 0.0, keep_with_next=True),
            "table_caption": ParagraphSpec("PSTableCaption", 10.5, False, "left", 1.5, 0.0, 6.0),
            "figure_title": ParagraphSpec("PSFigureTitle", 10.5, False, "center", 1.5, 0.0, 6.0, keep_with_next=True),
            "figure_note": ParagraphSpec("PSFigureNote", 10.5, False, "left", 1.5, 0.0, 2.0),
            "source_note": ParagraphSpec("PSSourceNote", 10.5, False, "left", 1.5, 0.0, 6.0),
            "figure_body": ParagraphSpec("PSFigureBody", 12.0, False, "center", 1.0, 0.0, 0.0),
        },
        table_font_size_pt=10.5,
        dense_table_font_size_pt=9.0,
        table_title_position="above",
        figure_title_position="below",
        table_caption_position="below",
        figure_note_position="below",
        source_note_position="below",
        page_margin_inch=0.95,
        figure_width_inch=6.10,
        page_width_inch=8.27,
        page_height_inch=11.69,
    )

    cn_kaiti_times = StyleProfile(
        name="cn_kaiti_times",
        font_profile=FontProfile(chinese="楷体", latin="Times New Roman"),
        role_specs=cn_song_times.role_specs,
        table_font_size_pt=cn_song_times.table_font_size_pt,
        dense_table_font_size_pt=cn_song_times.dense_table_font_size_pt,
        table_title_position=cn_song_times.table_title_position,
        figure_title_position=cn_song_times.figure_title_position,
        table_caption_position=cn_song_times.table_caption_position,
        figure_note_position=cn_song_times.figure_note_position,
        source_note_position=cn_song_times.source_note_position,
        page_margin_inch=cn_song_times.page_margin_inch,
        figure_width_inch=cn_song_times.figure_width_inch,
        page_width_inch=cn_song_times.page_width_inch,
        page_height_inch=cn_song_times.page_height_inch,
    )

    cn_heiti_arial = StyleProfile(
        name="cn_heiti_arial",
        font_profile=FontProfile(chinese="黑体", latin="Arial"),
        role_specs=cn_song_times.role_specs,
        table_font_size_pt=cn_song_times.table_font_size_pt,
        dense_table_font_size_pt=cn_song_times.dense_table_font_size_pt,
        table_title_position=cn_song_times.table_title_position,
        figure_title_position=cn_song_times.figure_title_position,
        table_caption_position=cn_song_times.table_caption_position,
        figure_note_position=cn_song_times.figure_note_position,
        source_note_position=cn_song_times.source_note_position,
        page_margin_inch=cn_song_times.page_margin_inch,
        figure_width_inch=cn_song_times.figure_width_inch,
        page_width_inch=cn_song_times.page_width_inch,
        page_height_inch=cn_song_times.page_height_inch,
    )

    teal_consulting_report = StyleProfile(
        name="teal_consulting_report",
        font_profile=FontProfile(chinese="微软雅黑", latin="Arial"),
        role_specs={
            "doc_title": ParagraphSpec("TealDocTitle", 26.0, True, "left", 1.0, 0.0, 12.0, 0.0, "#16845F", "Arial", "微软雅黑", True),
            "subtitle": ParagraphSpec("TealSubtitle", 11.0, False, "left", 1.1, 0.0, 10.0, 0.0, "#4D5559", "Arial", "微软雅黑", True),
            "heading_1": ParagraphSpec("TealHeading1", 28.0, True, "left", 1.05, 10.0, 14.0, 0.0, "#16845F", "Arial", "微软雅黑", True),
            "heading_2": ParagraphSpec("TealHeading2", 14.0, True, "left", 1.05, 10.0, 5.0, 0.0, "#30383D", "Arial", "微软雅黑", True),
            "heading_3": ParagraphSpec("TealHeading3", 10.0, True, "left", 1.05, 6.0, 2.0, 0.0, "#30383D", "Arial", "微软雅黑", True),
            "heading_4": ParagraphSpec("TealHeading4", 9.2, True, "left", 1.05, 5.0, 2.0, 0.0, "#30383D", "Arial", "微软雅黑", True),
            "lead": ParagraphSpec("TealLead", 10.0, False, "left", 1.12, 0.0, 8.0, 0.0, "#4D5559", "Arial", "微软雅黑"),
            "body": ParagraphSpec("TealBody", 9.2, False, "left", 1.08, 0.0, 6.0, 0.0, "#5F666A", "Arial", "微软雅黑"),
            "list": ParagraphSpec("TealList", 9.2, False, "left", 1.08, 0.0, 5.0, 0.0, "#5F666A", "Arial", "微软雅黑"),
            "table_title": ParagraphSpec("TealTableTitle", 10.0, True, "left", 1.0, 8.0, 3.0, 0.0, "#30383D", "Arial", "微软雅黑", True),
            "table_caption": ParagraphSpec("TealTableCaption", 7.0, False, "left", 1.0, 0.0, 6.0, 0.0, "#6F777B", "Arial", "微软雅黑"),
            "figure_title": ParagraphSpec("TealFigureTitle", 11.0, True, "left", 1.0, 8.0, 2.0, 0.0, "#30383D", "Arial", "微软雅黑", True),
            "figure_note": ParagraphSpec("TealFigureNote", 7.0, False, "left", 1.0, 0.0, 2.0, 0.0, "#6F777B", "Arial", "微软雅黑"),
            "source_note": ParagraphSpec("TealSourceNote", 7.0, False, "left", 1.0, 0.0, 6.0, 0.0, "#6F777B", "Arial", "微软雅黑"),
            "figure_body": ParagraphSpec("TealFigureBody", 9.2, False, "center", 1.0, 0.0, 0.0),
        },
        table_font_size_pt=9.0,
        dense_table_font_size_pt=8.5,
        table_title_position="above",
        figure_title_position="above",
        table_caption_position="below",
        figure_note_position="below",
        source_note_position="below",
        page_margin_inch=0.58,
        figure_width_inch=6.35,
        page_width_inch=8.27,
        page_height_inch=11.69,
        body_column_count=2,
        column_gap_inch=0.28,
    )

    red_private_equity_report = StyleProfile(
        name="red_private_equity_report",
        font_profile=FontProfile(chinese="宋体", latin="Times New Roman"),
        role_specs={
            "doc_title": ParagraphSpec("RedDocTitle", 18.0, True, "left", 1.0, 0.0, 10.0, 0.0, "#111111", "Arial", "微软雅黑"),
            "subtitle": ParagraphSpec("RedSubtitle", 9.5, False, "left", 1.1, 0.0, 8.0, 0.0, "#555555", "Arial", "微软雅黑"),
            "heading_1": ParagraphSpec("RedHeading1", 16.0, True, "left", 1.05, 12.0, 7.0, 0.0, "#111111", "Arial", "微软雅黑", True),
            "heading_2": ParagraphSpec("RedHeading2", 12.0, True, "left", 1.05, 10.0, 5.0, 0.0, "#111111", "Arial", "微软雅黑", True),
            "heading_3": ParagraphSpec("RedHeading3", 10.0, True, "left", 1.05, 6.0, 3.0, 0.0, "#111111", "Arial", "微软雅黑", True),
            "heading_4": ParagraphSpec("RedHeading4", 9.5, True, "left", 1.05, 5.0, 2.0, 0.0, "#111111", "Arial", "微软雅黑", True),
            "lead": ParagraphSpec("RedLead", 10.5, True, "left", 1.10, 0.0, 8.0, 0.0, "#222222", "Arial", "微软雅黑"),
            "body": ParagraphSpec("RedBody", 10.0, False, "left", 1.08, 0.0, 7.0),
            "list": ParagraphSpec("RedList", 9.5, False, "left", 1.08, 0.0, 6.0),
            "table_title": ParagraphSpec("RedTableTitle", 10.0, True, "left", 1.0, 10.0, 4.0, 0.0, "#111111", "Arial", "微软雅黑", True),
            "table_caption": ParagraphSpec("RedTableCaption", 6.5, False, "left", 1.0, 0.0, 6.0, 0.0, "#555555", "Arial", "微软雅黑"),
            "figure_title": ParagraphSpec("RedFigureTitle", 10.0, True, "left", 1.0, 10.0, 4.0, 0.0, "#111111", "Arial", "微软雅黑", True),
            "figure_note": ParagraphSpec("RedFigureNote", 6.5, False, "left", 1.0, 0.0, 2.0, 0.0, "#555555", "Arial", "微软雅黑"),
            "source_note": ParagraphSpec("RedSourceNote", 6.5, False, "left", 1.0, 0.0, 6.0, 0.0, "#555555", "Arial", "微软雅黑"),
            "figure_body": ParagraphSpec("RedFigureBody", 10.0, False, "center", 1.0, 0.0, 0.0),
        },
        table_font_size_pt=9.0,
        dense_table_font_size_pt=8.5,
        table_title_position="above",
        figure_title_position="above",
        table_caption_position="below",
        figure_note_position="below",
        source_note_position="below",
        page_margin_inch=0.78,
        figure_width_inch=6.25,
        page_width_inch=8.50,
        page_height_inch=11.0,
    )

    blue_editorial_article = StyleProfile(
        name="blue_editorial_article",
        font_profile=FontProfile(chinese="宋体", latin="Arial"),
        role_specs={
            "doc_title": ParagraphSpec("BlueDocTitle", 25.0, True, "left", 0.95, 0.0, 12.0, 0.0, "#071828", "Georgia", "宋体", True),
            "subtitle": ParagraphSpec("BlueSubtitle", 9.5, False, "left", 1.1, 0.0, 10.0, 0.0, "#5C6470", "Arial", "微软雅黑", True),
            "heading_1": ParagraphSpec("BlueHeading1", 15.5, True, "left", 1.02, 14.0, 7.0, 0.0, "#071828", "Georgia", "宋体", True),
            "heading_2": ParagraphSpec("BlueHeading2", 11.8, True, "left", 1.05, 10.0, 5.0, 0.0, "#071828", "Georgia", "宋体", True),
            "heading_3": ParagraphSpec("BlueHeading3", 10.0, True, "left", 1.08, 7.0, 3.0, 0.0, "#1E2C3A", "Arial", "微软雅黑", True),
            "heading_4": ParagraphSpec("BlueHeading4", 9.6, True, "left", 1.08, 6.0, 3.0, 0.0, "#1E2C3A", "Arial", "微软雅黑", True),
            "lead": ParagraphSpec("BlueLead", 11.0, True, "left", 1.28, 0.0, 10.0, 0.0, "#2E3A46", "Arial", "微软雅黑"),
            "body": ParagraphSpec("BlueBody", 10.0, False, "left", 1.24, 0.0, 8.0, 0.0, "#333333", "Arial", "微软雅黑"),
            "list": ParagraphSpec("BlueList", 9.5, False, "left", 1.18, 0.0, 6.0, 0.0, "#333333", "Arial", "微软雅黑"),
            "table_title": ParagraphSpec("BlueTableTitle", 9.5, True, "left", 1.0, 10.0, 3.0, 0.0, "#071828", "Arial", "微软雅黑", True),
            "table_caption": ParagraphSpec("BlueTableCaption", 6.5, False, "left", 1.0, 0.0, 7.0, 0.0, "#6A6F75", "Arial", "微软雅黑"),
            "figure_title": ParagraphSpec("BlueFigureTitle", 11.0, True, "left", 1.05, 12.0, 5.0, 0.0, "#071828", "Georgia", "宋体", True),
            "figure_note": ParagraphSpec("BlueFigureNote", 6.5, False, "left", 1.0, 0.0, 2.0, 0.0, "#6A6F75", "Arial", "微软雅黑"),
            "source_note": ParagraphSpec("BlueSourceNote", 6.5, False, "left", 1.0, 0.0, 8.0, 0.0, "#6A6F75", "Arial", "微软雅黑"),
            "figure_body": ParagraphSpec("BlueFigureBody", 10.0, False, "center", 1.0, 0.0, 0.0),
        },
        table_font_size_pt=8.8,
        dense_table_font_size_pt=8.0,
        table_title_position="above",
        figure_title_position="above",
        table_caption_position="below",
        figure_note_position="below",
        source_note_position="below",
        page_margin_inch=1.22,
        figure_width_inch=5.65,
        page_width_inch=8.50,
        page_height_inch=11.0,
    )

    return {
        cn_song_times.name: cn_song_times,
        cn_kaiti_times.name: cn_kaiti_times,
        cn_heiti_arial.name: cn_heiti_arial,
        teal_consulting_report.name: teal_consulting_report,
        red_private_equity_report.name: red_private_equity_report,
        blue_editorial_article.name: blue_editorial_article,
    }


STYLE_REGISTRY = build_style_registry()


def ensure_docx_support() -> None:
    """在真正进入 DOCX 构建或 QA 前确认 `python-docx` 可用。"""

    if DOCX_IMPORT_ERROR is not None:
        raise RuntimeError("当前环境缺少 `python-docx` 或相关依赖，无法执行 DOCX 路线。") from DOCX_IMPORT_ERROR


def get_style_profile(profile_name: str) -> StyleProfile:
    """根据名称返回 style profile。"""

    if profile_name not in STYLE_REGISTRY:
        raise KeyError(f"未知 style_profile: {profile_name}")
    return STYLE_REGISTRY[profile_name]


def resolve_meta_path(meta_path: Path) -> Path:
    """规范化 `meta.json` 路径。"""

    resolved = meta_path.resolve()
    if not resolved.exists():
        raise FileNotFoundError(f"meta.json 不存在: {resolved}")
    return resolved


def infer_workspace_root(meta_path: Path) -> Path:
    """从标准 refined 目录中推断 workspace 根目录。"""

    resolved = resolve_meta_path(meta_path)
    if len(resolved.parents) >= 3 and resolved.parents[1].name == "markdown":
        return resolved.parents[2]
    return resolved.parent


def load_meta(meta_path: Path) -> dict:
    """读取 `meta.json`。"""

    return json.loads(resolve_meta_path(meta_path).read_text(encoding="utf-8"))


def resolve_meta_declared_path(meta_path: Path, raw_path: str | None) -> Path | None:
    """把 `meta.json` 中声明的相对路径解析成绝对路径。"""

    if not raw_path:
        return None
    candidate = Path(raw_path)
    if candidate.is_absolute():
        return candidate
    workspace_root = infer_workspace_root(meta_path)
    return (workspace_root / candidate).resolve()


def resolve_standard_context(meta_path: Path) -> dict[str, Path | None]:
    """从标准 refined 元数据里解析常用路径。"""

    meta = load_meta(meta_path)
    markdown_path = resolve_meta_declared_path(meta_path, meta.get("markdown_file"))
    output_docx = resolve_meta_declared_path(meta_path, meta.get("output_docx"))
    assets_dir = resolve_meta_declared_path(meta_path, meta.get("assets_dir"))
    source_docx = resolve_meta_declared_path(meta_path, meta.get("source_docx"))
    asset_manifest_path = meta_path.parent / "asset_manifest.json"
    visual_review_path = infer_workspace_root(meta_path) / "temp" / "qa" / "visual_review.md"
    return {
        "workspace_root": infer_workspace_root(meta_path),
        "markdown_path": markdown_path,
        "output_docx": output_docx,
        "assets_dir": assets_dir,
        "source_docx": source_docx,
        "asset_manifest_path": asset_manifest_path if asset_manifest_path.exists() else None,
        "visual_review_path": visual_review_path if visual_review_path.exists() else None,
    }


def normalize_profile_name(profile_name: str) -> str:
    """规范化 style profile 名称。"""

    normalized = profile_name.strip()
    if normalized not in STYLE_REGISTRY:
        raise KeyError(f"未知 style_profile: {normalized}")
    return normalized


def build_markdown_template(title: str, profile_name: str, mode: str, include_figure_example: bool = False) -> str:
    """按模式和 profile 生成 Markdown 样板。"""

    if profile_name in ENGLISH_PRESET_PROFILES:
        lines = [
            f"# {title}",
            "",
            "Subtitle: Month YYYY | Fictional consulting report | Draft skeleton",
            "",
            "Lead: Replace this lead paragraph with the one-sentence claim that should anchor the document.",
            "",
            "## Executive summary",
            "",
            "Write the first body paragraph here.",
            "",
            "Table 1  Replace with the first table title",
            "| Column | Value |",
            "| --- | ---: |",
            "| Placeholder | 100 |",
            "Notes: Replace with the table note if needed.",
            "",
        ]
        if include_figure_example:
            lines.extend(
                [
                    "Exhibit 1  Replace with the first exhibit title",
                    "![Exhibit 1 placeholder](assets/placeholder.png)",
                    "Notes: Replace with the figure note.",
                    "Source: Replace with the source note.",
                    "",
                ]
            )
        if mode == "refined":
            lines.extend(
                [
                    "## Next section",
                    "",
                    "- Replace with the first governance point.",
                    "- Replace with the second governance point.",
                    "",
                ]
            )
        return "\n".join(lines) + "\n"

    lines = [
        f"# {title}",
        "",
        "副标题：日期 | 文档副标题 | Draft skeleton",
        "",
        "导语：把这段替换成一段正式导语。",
        "",
        "## 执行摘要",
        "",
        "把第一段正文写在这里。",
        "",
        "表 1 替换成第一张表的表题",
        "| 项目 | 数值 |",
        "| --- | ---: |",
        "| 占位项 | 100 |",
        "表注：按需替换表注。",
        "",
    ]
    if include_figure_example:
        lines.extend(
            [
                "![图 1 占位图](assets/placeholder.png)",
                "图 1 替换成第一张图的图题",
                "图注：按需替换图注。",
                "来源：按需替换来源说明。",
                "",
            ]
        )
    if mode == "refined":
        lines.extend(
            [
                "## 下一节",
                "",
                "- 替换成第一个要点。",
                "- 替换成第二个要点。",
                "",
            ]
        )
    return "\n".join(lines) + "\n"


def parse_markdown(markdown_text: str) -> list[Block]:
    """把受控 Markdown 文本解析成顺序稳定的 block 列表。"""

    lines = markdown_text.splitlines()
    index = 0
    blocks: list[Block] = []
    paragraph_buffer: list[str] = []
    seen_non_title_block = False

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer, seen_non_title_block
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer).strip()
        if text:
            blocks.append(Block(role="body", text=text))
            seen_non_title_block = True
        paragraph_buffer = []

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()

        if not line:
            flush_paragraph()
            index += 1
            continue

        image_match = IMAGE_PATTERN.fullmatch(line)
        if image_match:
            flush_paragraph()
            blocks.append(Block(role="figure_body", image_path=image_match.group("path")))
            seen_non_title_block = True
            index += 1
            continue

        if line.startswith("副标题：") or line.startswith("Subtitle:"):
            flush_paragraph()
            blocks.append(Block(role="subtitle", text=strip_known_prefix(line, ["副标题：", "Subtitle:"])))
            index += 1
            continue

        if line.startswith("导语：") or line.startswith("Lead:"):
            flush_paragraph()
            blocks.append(Block(role="lead", text=strip_known_prefix(line, ["导语：", "Lead:"])))
            seen_non_title_block = True
            index += 1
            continue

        if raw_line.startswith("#"):
            flush_paragraph()
            level = len(raw_line) - len(raw_line.lstrip("#"))
            title_text = raw_line[level:].strip()
            if level == 1 and not blocks and not seen_non_title_block:
                role = "doc_title"
            else:
                role = {
                    1: "heading_1",
                    2: "heading_1",
                    3: "heading_2",
                    4: "heading_3",
                }.get(level, "heading_4")
                seen_non_title_block = True
            blocks.append(Block(role=role, text=title_text))
            index += 1
            continue

        if line.startswith("- ") or ORDERED_LIST_PATTERN.match(line):
            flush_paragraph()
            while index < len(lines):
                current = lines[index].strip()
                if current.startswith("- "):
                    blocks.append(Block(role="list", text=f"• {current[2:].strip()}"))
                    seen_non_title_block = True
                    index += 1
                    continue
                if ORDERED_LIST_PATTERN.match(current):
                    blocks.append(Block(role="list", text=current))
                    seen_non_title_block = True
                    index += 1
                    continue
                break
            continue

        if line.startswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            blocks.append(Block(role="table_body", rows=parse_pipe_table(table_lines)))
            seen_non_title_block = True
            continue

        next_non_empty = peek_next_non_empty_line(lines, index + 1)
        previous_role = blocks[-1].role if blocks else None

        if line.startswith("表题："):
            flush_paragraph()
            blocks.append(Block(role="table_title", text=normalize_caption_text(line)))
            seen_non_title_block = True
            index += 1
            continue

        if line.startswith("图题："):
            flush_paragraph()
            blocks.append(Block(role="figure_title", text=normalize_caption_text(line)))
            seen_non_title_block = True
            index += 1
            continue

        if line.startswith("表注："):
            flush_paragraph()
            blocks.append(Block(role="table_caption", text=line))
            seen_non_title_block = True
            index += 1
            continue

        if line.startswith("图注："):
            flush_paragraph()
            blocks.append(Block(role="figure_note", text=line))
            seen_non_title_block = True
            index += 1
            continue

        if line.startswith("来源："):
            flush_paragraph()
            blocks.append(Block(role="source_note", text=line))
            seen_non_title_block = True
            index += 1
            continue

        if line.startswith("Note:") or line.startswith("Notes:"):
            flush_paragraph()
            if previous_role == "table_body":
                blocks.append(Block(role="table_caption", text=line))
            else:
                blocks.append(Block(role="figure_note", text=line))
            seen_non_title_block = True
            index += 1
            continue

        if line.startswith("Source:") or line.startswith("Sources:"):
            flush_paragraph()
            blocks.append(Block(role="source_note", text=line))
            seen_non_title_block = True
            index += 1
            continue

        if TABLE_TITLE_PATTERN.fullmatch(line) and next_non_empty.startswith("|"):
            flush_paragraph()
            blocks.append(Block(role="table_title", text=line))
            seen_non_title_block = True
            index += 1
            continue

        if FIGURE_TITLE_PATTERN.fullmatch(line) and (
            IMAGE_PATTERN.fullmatch(next_non_empty) is not None or previous_role == "figure_body"
        ):
            flush_paragraph()
            blocks.append(Block(role="figure_title", text=line))
            seen_non_title_block = True
            index += 1
            continue

        paragraph_buffer.append(raw_line)
        index += 1

    flush_paragraph()
    return blocks


def parse_pipe_table(lines: Iterable[str]) -> list[list[str]]:
    """解析 Markdown pipe table。"""

    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if not cells:
            continue
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    return rows


def strip_known_prefix(text: str, prefixes: list[str]) -> str:
    """移除文本命中的第一个已知前缀。"""

    normalized = text.strip()
    for prefix in prefixes:
        if normalized.startswith(prefix):
            return normalized.removeprefix(prefix).strip()
    return normalized


def normalize_caption_text(text: str) -> str:
    """去掉调试式 caption 前缀，保留正式文面。"""

    normalized = text.strip()
    for prefix in ("表题：", "图题："):
        if normalized.startswith(prefix):
            return normalized.removeprefix(prefix).strip()
    return normalized


def peek_next_non_empty_line(lines: list[str], start_index: int) -> str:
    """返回当前位置之后的下一条非空行。"""

    for index in range(start_index, len(lines)):
        if lines[index].strip():
            return lines[index].strip()
    return ""


def build_docx_from_markdown(markdown_path: Path, output_path: Path, profile_name: str) -> list[Block]:
    """读取 Markdown 并按指定 profile 构建 DOCX。"""

    ensure_docx_support()
    profile = get_style_profile(profile_name)
    markdown_text = markdown_path.read_text(encoding="utf-8")
    blocks = parse_markdown(markdown_text)

    document = Document()
    configure_document(document, profile)
    ensure_styles(document, profile)

    base_dir = markdown_path.parent
    current_section_mode = "single"
    heading_1_count = 0
    for block in blocks:
        if block.role == "heading_1":
            heading_1_count += 1
        target_section_mode = determine_section_mode(profile, block, heading_1_count)
        if target_section_mode != current_section_mode:
            switch_section_mode(document, profile, target_section_mode)
            current_section_mode = target_section_mode

        if block.role == "table_body":
            add_table(document, block.rows or [], profile)
            continue
        if block.role == "figure_body":
            add_figure(document, base_dir, block.image_path, profile)
            continue
        add_paragraph_block(document, block, profile)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return blocks


def configure_document(document: DocumentObject, profile: StyleProfile) -> None:
    """设置页面尺寸、页边距和初始 section。"""

    section = document.sections[0]
    section.page_width = Inches(profile.page_width_inch)
    section.page_height = Inches(profile.page_height_inch)
    apply_section_layout(section, profile, mode="single")
    section.start_type = WD_SECTION.NEW_PAGE


def determine_section_mode(profile: StyleProfile, block: Block, heading_1_count: int) -> str:
    """根据 profile 和 block 决定当前 section 的栏数模式。"""

    if profile.body_column_count <= 1:
        return "single"
    if block.role in {"doc_title", "subtitle", "lead"}:
        return "single"
    if block.role in {"table_title", "table_body", "table_caption", "figure_body", "figure_title", "figure_note", "source_note"}:
        return "single"
    if heading_1_count <= 1:
        return "single"
    return "double"


def apply_section_layout(section, profile: StyleProfile, mode: str) -> None:
    """统一设置 section 的页面参数和栏数。"""

    section.top_margin = Inches(profile.page_margin_inch)
    section.bottom_margin = Inches(profile.page_margin_inch)
    section.left_margin = Inches(profile.page_margin_inch)
    section.right_margin = Inches(profile.page_margin_inch)
    set_section_columns(
        section,
        column_count=profile.body_column_count if mode == "double" else 1,
        column_gap_inch=profile.column_gap_inch,
    )


def set_section_columns(section, column_count: int, column_gap_inch: float) -> None:
    """用 OOXML 直接设置 section 的分栏参数。"""

    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(column_count))
    cols.set(qn("w:space"), str(int(column_gap_inch * 1440)))


def switch_section_mode(document: DocumentObject, profile: StyleProfile, mode: str) -> None:
    """在文档末尾插入连续分节并切换栏数。"""

    section = document.add_section(WD_SECTION.CONTINUOUS)
    section.page_width = Inches(profile.page_width_inch)
    section.page_height = Inches(profile.page_height_inch)
    apply_section_layout(section, profile, mode)


def ensure_styles(document: DocumentObject, profile: StyleProfile) -> None:
    """在文档中创建全部 role 对应的样式。"""

    for spec in profile.role_specs.values():
        if spec.style_name in document.styles:
            continue
        style = document.styles.add_style(spec.style_name, WD_STYLE_TYPE.PARAGRAPH)
        configure_style(style, spec, profile)


def configure_style(style, spec: ParagraphSpec, profile: StyleProfile) -> None:
    """配置单个 Word 段落样式。"""

    font_profile = resolve_font_profile(profile.font_profile, spec)
    style.font.size = Pt(spec.size_pt)
    style.font.bold = spec.bold
    style.font.name = font_profile.latin
    if spec.color_hex:
        style.font.color.rgb = RGBColor.from_string(spec.color_hex.removeprefix("#"))
    patch_font_slots(style.element, font_profile)

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = spec.line_spacing
    paragraph_format.space_before = Pt(spec.space_before_pt)
    paragraph_format.space_after = Pt(spec.space_after_pt)
    paragraph_format.first_line_indent = Pt(spec.first_line_indent_pt)
    paragraph_format.keep_with_next = spec.keep_with_next


def add_paragraph_block(document: DocumentObject, block: Block, profile: StyleProfile) -> None:
    """把普通段落类 block 写入文档。"""

    spec = profile.role_specs[block.role]
    paragraph = document.add_paragraph(style=spec.style_name)
    paragraph.alignment = resolve_alignment(spec.alignment)
    paragraph.paragraph_format.line_spacing = spec.line_spacing
    paragraph.paragraph_format.space_before = Pt(spec.space_before_pt)
    paragraph.paragraph_format.space_after = Pt(spec.space_after_pt)
    paragraph.paragraph_format.first_line_indent = Pt(spec.first_line_indent_pt)
    paragraph.paragraph_format.keep_with_next = spec.keep_with_next
    add_inline_runs(paragraph, block.text, profile, spec)


def add_inline_runs(paragraph, text: str, profile: StyleProfile, spec: ParagraphSpec) -> None:
    """把支持 `**bold**` 的文本写入段落。"""

    segments = [segment for segment in INLINE_BOLD_PATTERN.split(text) if segment]
    if not segments:
        segments = [text]

    font_profile = resolve_font_profile(profile.font_profile, spec)
    for segment in segments:
        is_bold = segment.startswith("**") and segment.endswith("**")
        content = segment[2:-2] if is_bold else segment
        run = paragraph.add_run(content)
        run.bold = spec.bold or is_bold
        run.font.size = Pt(spec.size_pt)
        run.font.name = font_profile.latin
        if spec.color_hex:
            run.font.color.rgb = RGBColor.from_string(spec.color_hex.removeprefix("#"))
        patch_font_slots(run._element, font_profile)


def add_table(document: DocumentObject, rows: list[list[str]], profile: StyleProfile) -> None:
    """把 Markdown 表格写入文档，并按列角色设置对齐。"""

    if not rows:
        raise ValueError("表格内容为空，无法构建 DOCX。")

    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = True

    numeric_columns = detect_numeric_columns(rows)
    font_size = (
        profile.dense_table_font_size_pt
        if column_count >= 7 or any(len(cell) > 24 for row in rows for cell in row)
        else profile.table_font_size_pt
    )
    font_profile = profile.font_profile

    for row_index, row in enumerate(rows):
        if len(row) != column_count:
            raise ValueError("表格列数不一致，无法稳定构建 DOCX。")
        for col_index, cell_text in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = resolve_table_alignment(row_index, col_index, numeric_columns)

            run = paragraph.add_run(cell_text)
            run.bold = row_index == 0
            run.font.size = Pt(font_size)
            run.font.name = font_profile.latin
            patch_font_slots(run._element, font_profile)

            if profile.name == "blue_editorial_article":
                run.font.color.rgb = RGBColor.from_string(("071828" if row_index == 0 else "333333"))
                if row_index == 0:
                    set_cell_shading(cell, "E9EEF5")
            elif profile.name == "teal_consulting_report":
                run.font.color.rgb = RGBColor.from_string(("30383D" if row_index == 0 else "5F666A"))
                if row_index == 0:
                    set_cell_shading(cell, "EFF5F3")

    if profile.name == "blue_editorial_article":
        set_table_borders(table, "C9D3DF")
    elif profile.name == "teal_consulting_report":
        set_table_borders(table, "BFD6CE")


def add_figure(document: DocumentObject, base_dir: Path, image_path: str | None, profile: StyleProfile) -> None:
    """按 profile 规定的宽度插入图片。"""

    if not image_path:
        raise ValueError("图片 block 缺少路径。")
    resolved_path = (base_dir / image_path).resolve()
    if not resolved_path.exists():
        raise FileNotFoundError(f"图片不存在: {resolved_path}")

    spec = profile.role_specs["figure_body"]
    paragraph = document.add_paragraph(style=spec.style_name)
    paragraph.alignment = resolve_alignment(spec.alignment)
    paragraph.paragraph_format.space_before = Pt(spec.space_before_pt)
    paragraph.paragraph_format.space_after = Pt(spec.space_after_pt)
    run = paragraph.add_run()
    run.add_picture(str(resolved_path), width=Inches(profile.figure_width_inch))


def resolve_font_profile(base_font: FontProfile, spec: ParagraphSpec) -> FontProfile:
    """给某个 role 解析最终应使用的字体组合。"""

    return FontProfile(
        chinese=spec.chinese_font or base_font.chinese,
        latin=spec.latin_font or base_font.latin,
    )


def set_cell_shading(cell, fill_hex: str) -> None:
    """给单元格设置底色。"""

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def set_table_borders(table, color_hex: str) -> None:
    """统一设置表格边框颜色。"""

    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color_hex)


def patch_font_slots(target_element, font_profile: FontProfile) -> None:
    """显式设置 OOXML 字体槽位。"""

    rpr = target_element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        target_element.insert(0, rpr)

    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)

    rfonts.set(qn("w:ascii"), font_profile.latin)
    rfonts.set(qn("w:hAnsi"), font_profile.latin)
    rfonts.set(qn("w:eastAsia"), font_profile.chinese)
    rfonts.set(qn("w:cs"), font_profile.latin)


def resolve_alignment(alignment: str) -> int:
    """把字符串形式的对齐方式映射到 `python-docx` 枚举。"""

    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)


def detect_numeric_columns(rows: list[list[str]]) -> set[int]:
    """推断哪些列应按数值列右对齐。"""

    numeric_columns: set[int] = set()
    if len(rows) <= 1:
        return numeric_columns
    for col_index in range(1, len(rows[0])):
        values = [row[col_index].strip() for row in rows[1:] if col_index < len(row)]
        if values and all(is_numeric_text(value) for value in values):
            numeric_columns.add(col_index)
    return numeric_columns


def is_numeric_text(text: str) -> bool:
    """判断某个单元格文本是否可视为数值。"""

    return bool(NUMERIC_PATTERN.fullmatch(text.strip()))


def resolve_table_alignment(row_index: int, col_index: int, numeric_columns: set[int]) -> int:
    """根据表头、索引列和数值列确定对齐方式。"""

    if row_index == 0:
        return WD_ALIGN_PARAGRAPH.CENTER
    if col_index == 0:
        return WD_ALIGN_PARAGRAPH.LEFT
    if col_index in numeric_columns:
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def lint_markdown_source(
    markdown_path: Path,
    profile_name: str,
    workflow_mode: str,
    asset_manifest_path: Path | None = None,
) -> dict:
    """检查 Markdown 是否满足 skill 的语义契约。"""

    issues: list[Issue] = []
    if not markdown_path.exists():
        return {
            "passed": False,
            "issue_count": 1,
            "issues": [asdict(Issue("error", "markdown_missing", f"Markdown 不存在: {markdown_path}"))],
        }

    profile = get_style_profile(profile_name)
    markdown_text = markdown_path.read_text(encoding="utf-8")
    lines = markdown_text.splitlines()
    blocks = parse_markdown(markdown_text)
    markdown_dir = markdown_path.parent

    heading_levels: list[int] = []
    for line_number, raw_line in enumerate(lines, start=1):
        stripped = raw_line.strip()
        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            heading_levels.append(len(heading_match.group(1)))
        if stripped.startswith("表题：") or stripped.startswith("图题："):
            issues.append(Issue("warning", "deprecated_caption_prefix", f"第 {line_number} 行仍在使用调试式前缀：{stripped}"))
        if profile_name in ENGLISH_PRESET_PROFILES and (
            stripped.startswith("副标题：")
            or stripped.startswith("导语：")
            or stripped.startswith("表注：")
            or stripped.startswith("图注：")
            or stripped.startswith("来源：")
        ):
            issues.append(Issue("warning", "language_contract_mismatch", f"第 {line_number} 行使用了中文语义前缀，但 active profile 是英文 preset。"))

    for previous, current in zip(heading_levels, heading_levels[1:], strict=False):
        if current - previous > 1:
            issues.append(Issue("error", "heading_jump", f"标题层级从 H{previous} 直接跳到了 H{current}。"))

    if not blocks or blocks[0].role != "doc_title":
        issues.append(Issue("error", "missing_doc_title", "Markdown 没有以文档主标题开头。"))

    table_count = sum(1 for block in blocks if block.role == "table_body")
    table_title_count = sum(1 for block in blocks if block.role == "table_title")
    table_caption_count = sum(1 for block in blocks if block.role == "table_caption")
    figure_count = sum(1 for block in blocks if block.role == "figure_body")
    figure_title_count = sum(1 for block in blocks if block.role == "figure_title")
    figure_note_count = sum(1 for block in blocks if block.role == "figure_note")
    source_note_count = sum(1 for block in blocks if block.role == "source_note")

    if table_count and table_title_count != table_count:
        issues.append(Issue("error", "table_title_mismatch", "表格数量与表题数量不一致。"))
    if table_count and table_caption_count != table_count:
        issues.append(Issue("error", "table_caption_mismatch", "表格数量与表注数量不一致。"))
    if figure_count and figure_title_count != figure_count:
        issues.append(Issue("error", "figure_title_mismatch", "图片数量与图题数量不一致。"))

    if workflow_mode == "refined":
        if figure_count and figure_note_count != figure_count:
            issues.append(Issue("error", "figure_note_mismatch", "精细模式下图片数量与图注数量不一致。"))
        if figure_count and source_note_count != figure_count:
            issues.append(Issue("error", "source_note_mismatch", "精细模式下图片数量与来源说明数量不一致。"))
    else:
        if figure_count and figure_note_count != figure_count:
            issues.append(Issue("warning", "figure_note_missing", "轻量模式下部分图片缺少图注。"))
        if figure_count and source_note_count != figure_count:
            issues.append(Issue("warning", "source_note_missing", "轻量模式下部分图片缺少来源说明。"))

    missing_images: list[str] = []
    for block in blocks:
        if block.role == "figure_body":
            image_path = (markdown_dir / (block.image_path or "")).resolve()
            if not image_path.exists():
                missing_images.append(str(image_path))
    if missing_images:
        issues.append(Issue("error", "missing_image_asset", f"存在缺失图片资产: {missing_images}"))

    if asset_manifest_path is not None:
        if not asset_manifest_path.exists():
            issues.append(Issue("error", "asset_manifest_missing", f"asset_manifest 不存在: {asset_manifest_path}"))
        else:
            issues.extend(lint_asset_manifest(asset_manifest_path, markdown_dir, figure_count, profile))
    elif workflow_mode == "refined" and figure_count:
        issues.append(Issue("error", "asset_manifest_required", "精细模式包含复杂视觉资产，但缺少 asset_manifest.json。"))

    passed = not any(issue.severity == "error" for issue in issues)
    return {
        "passed": passed,
        "issue_count": len(issues),
        "issues": [asdict(issue) for issue in issues],
    }


def lint_asset_manifest(
    asset_manifest_path: Path,
    markdown_dir: Path,
    figure_count: int,
    profile: StyleProfile,
) -> list[Issue]:
    """检查 `asset_manifest.json` 的结构一致性。"""

    manifest = json.loads(asset_manifest_path.read_text(encoding="utf-8"))
    assets = manifest.get("assets")
    issues: list[Issue] = []
    if not isinstance(assets, list):
        return [Issue("error", "asset_manifest_invalid", "asset_manifest.json 缺少数组字段 `assets`。")]
    if figure_count and len(assets) != figure_count:
        issues.append(Issue("error", "asset_manifest_count_mismatch", "asset_manifest 中的资产数量与图片数量不一致。"))

    for asset in assets:
        asset_id = asset.get("asset_id", "<unknown>")
        asset_mode = asset.get("asset_mode")
        if not asset_mode:
            issues.append(Issue("error", "asset_mode_missing", f"{asset_id} 缺少 asset_mode。"))
        source_file = asset.get("source_file")
        if source_file and not (markdown_dir / source_file).exists():
            issues.append(Issue("error", "asset_source_missing", f"{asset_id} 的 source_file 不存在。"))
        caption_position = asset.get("caption_position")
        if caption_position and caption_position != profile.figure_title_position:
            issues.append(Issue("warning", "caption_policy_drift", f"{asset_id} 的 caption_position 与 active profile 不一致。"))
        if asset_mode == "python_figure" and not asset.get("generator_script"):
            issues.append(Issue("warning", "generator_script_missing", f"{asset_id} 缺少 generator_script。"))
        if not asset.get("figure_note"):
            issues.append(Issue("warning", "figure_note_missing", f"{asset_id} 缺少 figure_note。"))
        if not asset.get("source_note"):
            issues.append(Issue("warning", "source_note_missing", f"{asset_id} 缺少 source_note。"))
    return issues


def run_docx_qa(
    markdown_path: Path,
    docx_path: Path,
    profile_name: str,
    workflow_mode: str,
    meta_path: Path | None = None,
    asset_manifest_path: Path | None = None,
    visual_review_path: Path | None = None,
    require_visual_review: bool = True,
) -> dict:
    """执行 DOCX 自动 QA，并返回可落盘的结构化报告。"""

    ensure_docx_support()
    profile = get_style_profile(profile_name)
    blocks = parse_markdown(markdown_path.read_text(encoding="utf-8"))
    document = Document(docx_path)

    results = {
        "source_integrity": check_source_integrity(
            markdown_path=markdown_path,
            docx_path=docx_path,
            meta_path=meta_path,
            asset_manifest_path=asset_manifest_path,
            blocks=blocks,
            profile=profile,
            workflow_mode=workflow_mode,
        ),
        "style_contract": check_style_contract(document, profile),
        "font_slot_integrity": check_font_slots(docx_path),
        "block_sequence": check_block_sequence(docx_path, blocks),
        "section_contract": check_section_contract(docx_path, profile),
        "asset_manifest_integrity": check_asset_manifest(asset_manifest_path, markdown_path.parent, blocks, profile, workflow_mode),
        "visual_review_status": check_visual_review_record(visual_review_path, require_visual_review),
    }

    passed_all_auto = all(result.passed for name, result in results.items() if name != "visual_review_status")
    passed_all = all(result.passed for result in results.values())
    return {
        "markdown_path": str(markdown_path),
        "docx_path": str(docx_path),
        "meta_path": str(meta_path) if meta_path else None,
        "asset_manifest_path": str(asset_manifest_path) if asset_manifest_path else None,
        "visual_review_path": str(visual_review_path) if visual_review_path else None,
        "style_profile": profile.name,
        "workflow_mode": workflow_mode,
        "passed_all_auto_checks": passed_all_auto,
        "passed_all_checks": passed_all,
        "checks": {name: asdict(result) for name, result in results.items()},
    }


def check_source_integrity(
    markdown_path: Path,
    docx_path: Path,
    meta_path: Path | None,
    asset_manifest_path: Path | None,
    blocks: list[Block],
    profile: StyleProfile,
    workflow_mode: str,
) -> CheckResult:
    """检查 Markdown、DOCX、meta 和图片资产是否完整。"""

    if not markdown_path.exists():
        return CheckResult(False, f"Markdown 不存在: {markdown_path}")
    if not docx_path.exists():
        return CheckResult(False, f"DOCX 不存在: {docx_path}")
    if workflow_mode == "refined" and meta_path is not None and not meta_path.exists():
        return CheckResult(False, f"meta.json 不存在: {meta_path}")

    if meta_path is not None:
        meta = load_meta(meta_path)
        if meta.get("style_profile") and meta.get("style_profile") != profile.name:
            return CheckResult(False, f"meta.json 的 style_profile 不是 {profile.name}。")
        if meta.get("workflow_mode") and meta.get("workflow_mode") != workflow_mode:
            return CheckResult(False, f"meta.json 的 workflow_mode 不是 {workflow_mode}。")

    base_dir = markdown_path.parent
    missing_images: list[str] = []
    for block in blocks:
        if block.role == "figure_body":
            image_path = (base_dir / (block.image_path or "")).resolve()
            if not image_path.exists():
                missing_images.append(str(image_path))
    if missing_images:
        return CheckResult(False, f"存在缺失图片资产: {missing_images}")
    if workflow_mode == "refined" and any(block.role == "figure_body" for block in blocks) and asset_manifest_path is None:
        return CheckResult(False, "精细模式存在复杂视觉资产，但未提供 asset_manifest.json 路径。")
    return CheckResult(True, "Markdown、DOCX 与图片资产完整，基础 source 契约通过。")


def check_style_contract(document: Document, profile: StyleProfile) -> CheckResult:
    """检查段落、缩进、表题加粗和表格对齐是否满足契约。"""

    style_specs = {spec.style_name: spec for spec in profile.role_specs.values()}
    failures: list[str] = []
    paragraph_samples = 0

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name
        spec = style_specs.get(style_name)
        if spec is None:
            continue
        text = paragraph.text.strip()
        if not text and not paragraph.runs:
            continue
        paragraph_samples += 1

        first_run = first_non_empty_run(paragraph.runs)
        if first_run is not None:
            run_size = _pt_value(first_run.font.size) or spec.size_pt
            if not is_close(run_size, spec.size_pt, 0.35):
                failures.append(f"{style_name} 字号为 {run_size}pt，预期 {spec.size_pt}pt。")
            if spec.bold and first_run.bold is not True:
                failures.append(f"{style_name} 首个文本 run 未加粗。")

        alignment = paragraph.alignment
        expected_alignment = resolve_alignment(spec.alignment)
        if alignment != expected_alignment:
            failures.append(f"{style_name} 对齐方式不符合预期。")

        line_spacing = paragraph.paragraph_format.line_spacing or paragraph.style.paragraph_format.line_spacing
        if not is_close(float(line_spacing), spec.line_spacing, 0.02):
            failures.append(f"{style_name} 行距为 {line_spacing}，预期 {spec.line_spacing}。")

        space_before = _pt_value(paragraph.paragraph_format.space_before)
        if space_before is None:
            space_before = _pt_value(paragraph.style.paragraph_format.space_before) or 0.0
        if not is_close(space_before, spec.space_before_pt, 0.25):
            failures.append(f"{style_name} 段前为 {space_before}pt，预期 {spec.space_before_pt}pt。")

        space_after = _pt_value(paragraph.paragraph_format.space_after)
        if space_after is None:
            space_after = _pt_value(paragraph.style.paragraph_format.space_after) or 0.0
        if not is_close(space_after, spec.space_after_pt, 0.25):
            failures.append(f"{style_name} 段后为 {space_after}pt，预期 {spec.space_after_pt}pt。")

        first_line_indent = _pt_value(paragraph.paragraph_format.first_line_indent)
        if first_line_indent is None:
            first_line_indent = _pt_value(paragraph.style.paragraph_format.first_line_indent) or 0.0
        if not is_close(first_line_indent, spec.first_line_indent_pt, 0.25):
            failures.append(f"{style_name} 首行缩进为 {first_line_indent}pt，预期 {spec.first_line_indent_pt}pt。")

    failures.extend(check_table_contract(document, profile))
    if paragraph_samples == 0:
        failures.append("没有找到任何可检查的段落样本。")
    if failures:
        return CheckResult(False, " ; ".join(failures))
    return CheckResult(True, f"段落、缩进、表题加粗和表格对齐契约通过，共检查 {paragraph_samples} 个段落样本。")


def check_table_contract(document: Document, profile: StyleProfile) -> list[str]:
    """检查表格字号和列对齐。"""

    failures: list[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        numeric_columns = detect_numeric_columns(rows)
        allowed_sizes = {profile.table_font_size_pt, profile.dense_table_font_size_pt}

        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                paragraph = cell.paragraphs[0]
                expected_alignment = resolve_table_alignment(row_index, col_index, numeric_columns)
                if paragraph.alignment != expected_alignment:
                    failures.append(f"表 {table_index} 第 {row_index + 1} 行第 {col_index + 1} 列对齐方式不符合预期。")
                    return failures

                first_run = first_non_empty_run(paragraph.runs)
                if first_run is None:
                    continue
                size_pt = _pt_value(first_run.font.size) or 0.0
                if not any(is_close(size_pt, allowed_size, 0.4) for allowed_size in allowed_sizes):
                    failures.append(f"表 {table_index} 第 {row_index + 1} 行第 {col_index + 1} 列字号为 {size_pt}pt，不在允许范围内。")
                    return failures
                if row_index == 0 and first_run.bold is not True:
                    failures.append(f"表 {table_index} 表头未加粗。")
                    return failures
    return failures


def check_font_slots(docx_path: Path) -> CheckResult:
    """检查文档里文本 run 的字体槽位是否完整。"""

    with ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml")

    root = ET.fromstring(document_xml)
    run_count = 0
    missing_count = 0
    for run in root.findall(".//w:r", W_NAMESPACE):
        text_node = run.find("w:t", W_NAMESPACE)
        if text_node is None or not (text_node.text or "").strip():
            continue
        run_count += 1
        rfonts = run.find("w:rPr/w:rFonts", W_NAMESPACE)
        if rfonts is None:
            missing_count += 1
            continue
        attrs = rfonts.attrib
        required = [
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii",
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi",
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs",
        ]
        if any(required_attr not in attrs for required_attr in required):
            missing_count += 1

    if run_count == 0:
        return CheckResult(False, "文档里没有检测到可校验的文本 run。")
    if missing_count:
        return CheckResult(False, f"{missing_count}/{run_count} 个文本 run 缺少完整字体槽位。")
    return CheckResult(True, f"全部 {run_count} 个文本 run 都具备完整字体槽位。")


def check_block_sequence(docx_path: Path, expected_blocks: list[Block]) -> CheckResult:
    """检查文档 block 顺序与 Markdown 语义顺序是否一致。"""

    with ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml")

    root = ET.fromstring(document_xml)
    actual_sequence: list[str] = []
    body = root.find("w:body", W_NAMESPACE)
    if body is None:
        return CheckResult(False, "DOCX 中没有 w:body。")

    for child in body:
        tag = child.tag.rsplit("}", maxsplit=1)[-1]
        if tag == "tbl":
            actual_sequence.append("table_body")
            continue
        if tag != "p":
            continue
        style = child.find("w:pPr/w:pStyle", W_NAMESPACE)
        style_value = style.attrib.get(W_VAL) if style is not None else ""
        if child.find(".//w:drawing", W_NAMESPACE) is not None:
            actual_sequence.append("figure_body")
            continue
        role = STYLE_ID_TO_ROLE.get(style_value)
        if role:
            actual_sequence.append(role)

    expected_sequence = [block.role for block in expected_blocks]
    if expected_sequence != actual_sequence:
        return CheckResult(False, f"文档 block 顺序与 Markdown 不一致。期望 {expected_sequence}，实际 {actual_sequence}。")
    return CheckResult(True, f"文档 block 顺序与 Markdown 一致，共检查 {len(expected_sequence)} 个 block。")


def check_section_contract(docx_path: Path, profile: StyleProfile) -> CheckResult:
    """检查 section 栏数是否符合 profile 约定。"""

    section_columns = extract_section_columns(docx_path)
    if not section_columns:
        return CheckResult(False, "没有在 DOCX 包中找到任何 section 定义。")
    if profile.body_column_count == 1:
        if any(column_count != 1 for column_count in section_columns):
            return CheckResult(False, f"当前 profile 预期单栏，但 section 栏数为 {section_columns}。")
        return CheckResult(True, f"section 栏数契约通过，全部为单栏：{section_columns}。")

    if profile.body_column_count not in section_columns:
        return CheckResult(False, f"当前 profile 预期出现 {profile.body_column_count} 栏 section，但实际只有 {section_columns}。")
    if 1 not in section_columns:
        return CheckResult(False, f"当前 profile 需要标题/图表的单栏 section，但实际栏数为 {section_columns}。")
    return CheckResult(True, f"section 栏数契约通过，实际栏数序列为 {section_columns}。")


def extract_section_columns(docx_path: Path) -> list[int]:
    """从 DOCX XML 中提取每个 section 的栏数。"""

    with ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml")
    root = ET.fromstring(document_xml)

    section_columns: list[int] = []
    for sect_pr in root.findall(".//w:sectPr", W_NAMESPACE):
        cols = sect_pr.find("w:cols", W_NAMESPACE)
        if cols is None:
            section_columns.append(1)
            continue
        raw_value = cols.attrib.get(W_VAL) or cols.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num")
        section_columns.append(int(raw_value) if raw_value else 1)
    return section_columns


def check_asset_manifest(
    asset_manifest_path: Path | None,
    markdown_dir: Path,
    blocks: list[Block],
    profile: StyleProfile,
    workflow_mode: str,
) -> CheckResult:
    """检查 asset manifest 与图表资产声明是否一致。"""

    figure_count = sum(1 for block in blocks if block.role == "figure_body")
    if figure_count == 0:
        return CheckResult(True, "文档没有复杂视觉资产，因此不要求 asset manifest。")
    if workflow_mode != "refined":
        return CheckResult(True, "当前不是精细模式，因此跳过 asset manifest 强制检查。")
    if asset_manifest_path is None or not asset_manifest_path.exists():
        return CheckResult(False, "精细模式存在复杂视觉资产，但 asset_manifest.json 不存在。")

    issues = lint_asset_manifest(asset_manifest_path, markdown_dir, figure_count, profile)
    errors = [issue.message for issue in issues if issue.severity == "error"]
    if errors:
        return CheckResult(False, " ; ".join(errors))
    return CheckResult(True, f"asset_manifest 与 {figure_count} 个图表资产声明一致。")


def check_visual_review_record(visual_review_path: Path | None, require_visual_review: bool) -> CheckResult:
    """检查是否已有人工或证据驱动的 visual review 记录。"""

    if not require_visual_review:
        return CheckResult(True, "当前未强制要求 visual review。")
    if visual_review_path is None:
        return CheckResult(False, "未提供 visual review 记录路径。")
    if not visual_review_path.exists():
        return CheckResult(False, f"visual review 记录不存在: {visual_review_path}")
    review_text = visual_review_path.read_text(encoding="utf-8").strip()
    if len(review_text) < 40 or "待补充" in review_text:
        return CheckResult(False, "visual review 记录过短，或仍处于待补充状态。")
    return CheckResult(True, f"已发现 visual review 记录: {visual_review_path}")


def export_docx_preview(
    docx_path: Path,
    preview_dir: Path,
    png_prefix: str = "page",
    dpi: int = DEFAULT_PDF_DPI,
    export_png: bool = True,
) -> dict:
    """把 DOCX 导出为 PDF，并按页切成 PNG。"""

    office_binary = shutil.which("soffice") or shutil.which("libreoffice")
    if office_binary is None:
        raise FileNotFoundError("当前环境缺少 soffice / libreoffice，无法导出 DOCX 预览。")

    preview_dir.mkdir(parents=True, exist_ok=True)
    pages_dir = preview_dir / "pages"
    pages_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = preview_dir / f"{docx_path.stem}.pdf"

    for stale_png in pages_dir.glob(f"{png_prefix}-*.png"):
        stale_png.unlink()

    subprocess.run(
        [
            office_binary,
            "--headless",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(preview_dir),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    generated_pdf = preview_dir / f"{docx_path.stem}.pdf"
    if generated_pdf != pdf_path and generated_pdf.exists():
        generated_pdf.replace(pdf_path)

    if export_png:
        pdftoppm = shutil.which("pdftoppm")
        if pdftoppm is None:
            raise FileNotFoundError("当前环境缺少 pdftoppm，无法导出逐页 PNG。")
        subprocess.run(
            [
                pdftoppm,
                "-png",
                "-r",
                str(dpi),
                str(pdf_path),
                str(pages_dir / png_prefix),
            ],
            check=True,
            capture_output=True,
            text=True,
        )

    page_images = sorted(pages_dir.glob(f"{png_prefix}-*.png"))
    return {
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path),
        "preview_dir": str(preview_dir),
        "pages_dir": str(pages_dir),
        "page_image_count": len(page_images),
    }


def guess_preview_dir(docx_path: Path) -> Path:
    """按标准 `build/docx/*.docx` 结构推断 preview 目录。"""

    resolved = docx_path.resolve()
    if len(resolved.parents) >= 3 and resolved.parents[0].name == "docx" and resolved.parents[1].name == "build":
        return resolved.parents[2] / "temp" / "preview"
    return resolved.parent / "preview"


def guess_qa_dir(docx_path: Path) -> Path:
    """按标准 `build/docx/*.docx` 结构推断 QA 目录。"""

    resolved = docx_path.resolve()
    if len(resolved.parents) >= 3 and resolved.parents[0].name == "docx" and resolved.parents[1].name == "build":
        return resolved.parents[2] / "temp" / "qa"
    return resolved.parent / "qa"


def first_non_empty_run(runs):
    """返回段落里第一个包含实际文本的 run。"""

    for run in runs:
        if (run.text or "").strip():
            return run
    return None


def _pt_value(length) -> float | None:
    """把 `python-docx` 长度对象稳定转换成 pt。"""

    if length is None:
        return None
    return round(length.pt, 1)


def is_close(actual: float, expected: float, tolerance: float) -> bool:
    """判断两个数值是否在允许误差范围内。"""

    return abs(actual - expected) <= tolerance


def report_to_markdown(title: str, report: dict) -> str:
    """把 lint 或 QA 报告稳定渲染成 Markdown。"""

    lines = [f"# {title}", ""]
    if "issues" in report:
        lines.append(f"- passed: `{report['passed']}`")
        lines.append(f"- issue_count: `{report['issue_count']}`")
        lines.append("")
        for issue in report["issues"]:
            lines.append(f"## {issue['code']}")
            lines.append("")
            lines.append(f"- severity: `{issue['severity']}`")
            lines.append(f"- message: {issue['message']}")
            lines.append("")
        if not report["issues"]:
            lines.append("- 没有发现问题。")
            lines.append("")
        return "\n".join(lines)

    lines.append(f"- 自动检查是否全部通过：`{report['passed_all_auto_checks']}`")
    lines.append(f"- 全部检查是否通过：`{report['passed_all_checks']}`")
    lines.append(f"- style profile：`{report['style_profile']}`")
    lines.append(f"- workflow mode：`{report['workflow_mode']}`")
    lines.append(f"- Markdown：`{report['markdown_path']}`")
    lines.append(f"- DOCX：`{report['docx_path']}`")
    lines.append(f"- Asset manifest：`{report['asset_manifest_path']}`")
    lines.append(f"- Visual review：`{report['visual_review_path']}`")
    lines.append("")
    for check_name, check_result in report["checks"].items():
        lines.append(f"## {check_name}")
        lines.append("")
        lines.append(f"- passed: `{check_result['passed']}`")
        lines.append(f"- detail: {check_result['detail']}")
        lines.append("")
    return "\n".join(lines)


def write_json(path: Path, payload: dict) -> None:
    """把结构化结果稳定写成 JSON。"""

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_markdown(path: Path, title: str, payload: dict) -> None:
    """把 lint 或 QA 报告写成 Markdown。"""

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(report_to_markdown(title, payload) + "\n", encoding="utf-8")
